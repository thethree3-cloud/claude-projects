"""Résumé file -> plain text, so the UI can take an upload instead of a paste.

`parse_resume` only speaks text; this module is the thin adapter in front of
it. PDF goes through PyMuPDF (already a repo dependency, used by Projects
01/04/14). .docx is a zip of XML — python-docx is used when it's installed,
otherwise a small built-in reader pulls the paragraph text straight out of
`word/document.xml` (no extra dependency). .txt/.md are decoded as-is.

    text = extract_text(uploaded.read(), uploaded.name)
    resume = parse_resume(text)
"""

import html
import io
import re
import zipfile

SUPPORTED = (".pdf", ".docx", ".txt", ".md", ".markdown")


def _from_pdf(data: bytes) -> str:
    import fitz  # PyMuPDF

    with fitz.open(stream=data, filetype="pdf") as doc:
        return "\n".join(page.get_text() for page in doc)


def _from_docx(data: bytes) -> str:
    try:
        import docx  # python-docx, if available
    except ModuleNotFoundError:
        return _docx_without_lib(data)

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _docx_without_lib(data: bytes) -> str:
    """Minimal .docx reader: the body text lives in `word/document.xml` as
    <w:t> runs grouped into <w:p> paragraphs. Good enough for a résumé; it
    won't reconstruct complex tables."""
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
    xml = re.sub(r"<w:tab\b[^>]*/>", "\t", xml)
    xml = re.sub(r"</w:p>", "\n", xml)
    text = re.sub(r"<[^>]+>", "", xml)
    return html.unescape(text)


def extract_text(data: bytes, filename: str) -> str:
    """Return the plain text of a résumé file. `filename` is used only for its
    extension. Raises ValueError for an unsupported type and RuntimeError if
    the file is the right type but can't be read."""
    name = (filename or "").lower()
    ext = name[name.rfind(".") :] if "." in name else ""
    if ext not in SUPPORTED:
        raise ValueError(
            f"Unsupported résumé file '{filename}'. Use one of: "
            + ", ".join(SUPPORTED)
        )
    try:
        if ext == ".pdf":
            text = _from_pdf(data)
        elif ext == ".docx":
            text = _from_docx(data)
        else:
            text = data.decode("utf-8", errors="replace")
    except (ValueError, RuntimeError):
        raise
    except Exception as exc:  # noqa: BLE001 — corrupt file, wrong lib, etc.
        raise RuntimeError(f"Couldn't read '{filename}': {exc}") from exc

    text = text.replace("\r\n", "\n").strip()
    if not text:
        raise RuntimeError(
            f"'{filename}' had no extractable text — a scanned/image PDF?"
        )
    return text
